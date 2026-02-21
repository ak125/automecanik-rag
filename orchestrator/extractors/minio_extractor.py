"""MinIO Extractor - Extract documents from S3-compatible storage.

Routes files to specialized pipelines based on extension:
- PDF  -> ingest_pdfs pipeline (H2/H3-aware chunking + OCR)
- Images -> ingest_images pipeline (OCR + vision captions)
- Logs -> ingest_logs pipeline (profile + error patterns)
- CSV/XLSX -> tabular pipeline
- DOCX/Markdown/Text -> inline extraction

The extractor downloads files from MinIO buckets, converts them to markdown,
and returns them for chunking and indexing.
"""

import logging
from typing import List, Optional
from pathlib import Path
import tempfile
import os

logger = logging.getLogger(__name__)

# Format routing map: extension -> pipeline type
FORMAT_HANDLERS = {
    ".pdf": "pdf",
    ".csv": "tabular",
    ".xlsx": "tabular",
    ".xls": "tabular",
    ".json": "structured",
    ".xml": "structured",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
    ".bmp": "image",
    ".tif": "image",
    ".tiff": "image",
    ".log": "log",
    ".out": "log",
    ".err": "log",
    ".txt": "text",
    ".md": "markdown",
    ".docx": "docx",
}


class MinIOExtractor:
    """Extract documents from MinIO S3-compatible storage."""

    def __init__(
        self,
        endpoint: str = "minio:9000",
        access_key: Optional[str] = None,
        secret_key: Optional[str] = None,
        buckets: Optional[List[str]] = None,
        secure: bool = False,
    ):
        """Initialize MinIO extractor.

        Args:
            endpoint: MinIO server endpoint
            access_key: Access key (or from MINIO_ACCESS_KEY env)
            secret_key: Secret key (or from MINIO_SECRET_KEY env)
            buckets: List of buckets to extract from
            secure: Use HTTPS
        """
        self.endpoint = endpoint
        self.access_key = access_key or os.getenv("MINIO_ACCESS_KEY", "")
        self.secret_key = secret_key or os.getenv("MINIO_SECRET_KEY", "")
        self.buckets = buckets or ["docs", "exports"]
        self.secure = secure
        self._client = None
        self._temp_dir = tempfile.mkdtemp(prefix="minio_")

    def _get_client(self):
        """Get or create MinIO client."""
        if self._client is None:
            try:
                from minio import Minio

                self._client = Minio(
                    self.endpoint,
                    access_key=self.access_key,
                    secret_key=self.secret_key,
                    secure=self.secure,
                )
            except ImportError:
                logger.error("minio package not installed")
                return None
        return self._client

    async def extract(self) -> List[dict]:
        """Extract all documents from configured buckets.

        Returns:
            List of document dicts with content and metadata.
        """
        docs = []
        client = self._get_client()

        if client is None:
            return docs

        for bucket_name in self.buckets:
            bucket_docs = await self._extract_bucket(client, bucket_name)
            docs.extend(bucket_docs)

        return docs

    async def _extract_bucket(self, client, bucket_name: str) -> List[dict]:
        """Extract documents from a single bucket."""
        docs = []

        try:
            # Check if bucket exists
            if not client.bucket_exists(bucket_name):
                logger.warning(f"Bucket not found: {bucket_name}")
                return docs

            # List all objects
            objects = client.list_objects(bucket_name, recursive=True)

            for obj in objects:
                try:
                    doc = await self._extract_object(client, bucket_name, obj.object_name)
                    if doc:
                        docs.append(doc)
                except Exception as e:
                    logger.error(f"Failed to extract {bucket_name}/{obj.object_name}: {e}")

        except Exception as e:
            logger.error(f"Failed to list bucket {bucket_name}: {e}")

        return docs

    async def _extract_object(self, client, bucket_name: str, object_name: str) -> Optional[dict]:
        """Extract content from a single object, routing to specialized pipelines."""
        ext = Path(object_name).suffix.lower()
        format_type = FORMAT_HANDLERS.get(ext)

        if format_type is None:
            logger.debug(f"Skipping unsupported file type: {object_name}")
            return None

        # Download to temp file
        local_path = Path(self._temp_dir) / Path(object_name).name
        try:
            client.fget_object(bucket_name, object_name, str(local_path))
        except Exception as e:
            logger.error(f"Failed to download {object_name}: {e}")
            return None

        content = None
        source_type = "minio"
        try:
            if format_type == "pdf":
                content, source_type = self._extract_via_pdf_pipeline(local_path)
            elif format_type == "image":
                content, source_type = self._extract_via_image_pipeline(local_path)
            elif format_type == "log":
                content, source_type = self._extract_via_log_pipeline(local_path, object_name)
            elif format_type in ("markdown", "text"):
                content = local_path.read_text(encoding="utf-8", errors="replace")
            elif format_type == "docx":
                content = self._extract_docx(local_path)
            else:
                # tabular, structured, etc. — read as text fallback
                content = local_path.read_text(encoding="utf-8", errors="replace")

            logger.info(f"Routed {object_name} to {format_type} pipeline")
        finally:
            if local_path.exists():
                local_path.unlink()

        if not content:
            return None

        return {
            "source_path": f"minio/{bucket_name}/{object_name}",
            "content": content,
            "title": Path(object_name).stem,
            "source_type": source_type,
            "metadata": {
                "bucket": bucket_name,
                "object_name": object_name,
                "file_type": ext,
                "pipeline": format_type,
            },
        }

    def _extract_via_pdf_pipeline(self, file_path: Path) -> tuple[Optional[str], str]:
        """Route PDF to specialized pipeline with page markers and OCR."""
        try:
            # Import from ingest_pdfs for full extraction (page markers + OCR)
            from scripts.ingestors.ingest_pdfs import extract_pdf_text
            content, stats = extract_pdf_text(file_path, enable_ocr=True)
            logger.debug(f"PDF pipeline: pages_text={stats['pages_text']} pages_ocr={stats['pages_ocr']}")
            return content, "guide"
        except ImportError:
            # Fallback to basic extraction
            return self._extract_pdf(file_path), "minio"

    def _extract_via_image_pipeline(self, file_path: Path) -> tuple[Optional[str], str]:
        """Route image to OCR + caption pipeline."""
        try:
            from scripts.ingestors.ingest_images import extract_ocr_blocks, build_caption, render_ocr_markdown
            ocr_blocks = extract_ocr_blocks(file_path)
            caption = build_caption(file_path, ocr_blocks)
            ocr_md = render_ocr_markdown(ocr_blocks)
            content = f"## Caption\n\n{caption}\n\n{ocr_md}"
            return content, "media"
        except ImportError:
            logger.warning("ingest_images not available, skipping image")
            return None, "media"

    def _extract_via_log_pipeline(self, file_path: Path, object_name: str) -> tuple[Optional[str], str]:
        """Route log to profile + error pattern pipeline."""
        try:
            from scripts.ingestors.ingest_logs import extract_log_profile, render_profile_markdown
            content = file_path.read_text(encoding="utf-8", errors="replace")
            profile = extract_log_profile(content, object_name)
            md = render_profile_markdown(profile)
            return md, "diagnostic"
        except ImportError:
            logger.warning("ingest_logs not available, skipping log")
            return None, "diagnostic"

    def _extract_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            # Try pypdf first
            try:
                from pypdf import PdfReader

                reader = PdfReader(str(file_path))
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
            except ImportError:
                pass

            # Fallback to pdfplumber
            try:
                import pdfplumber

                with pdfplumber.open(str(file_path)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return "\n\n".join(text_parts)
            except ImportError:
                pass

            logger.warning("No PDF extraction library available (pypdf or pdfplumber)")
            return None

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return None

    def _extract_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            from docx import Document

            doc = Document(str(file_path))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except ImportError:
            logger.warning("python-docx not installed")
            return None
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None

    def cleanup(self):
        """Cleanup temp directory."""
        import shutil

        if Path(self._temp_dir).exists():
            shutil.rmtree(self._temp_dir)


async def extract_minio_documents(config) -> List:
    """Extract documents from MinIO using config.

    Args:
        config: PipelineConfig with MinIO settings.

    Returns:
        List of Document objects.
    """
    from ..pipeline import Document

    extractor = MinIOExtractor(
        endpoint=config.minio_endpoint,
        access_key=config.minio_access_key,
        secret_key=config.minio_secret_key,
        buckets=config.minio_buckets,
    )

    try:
        raw_docs = await extractor.extract()
        logger.info(f"Extracted {len(raw_docs)} documents from MinIO")

        return [
            Document(
                source_path=doc["source_path"],
                content=doc["content"],
                title=doc["title"],
                source_type="minio",
                truth_level="L3",  # MinIO docs default to L3 (Hypotheses)
                metadata=doc.get("metadata", {}),
            )
            for doc in raw_docs
        ]
    finally:
        extractor.cleanup()
